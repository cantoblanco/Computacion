import sys
import os
from PyQt6.QtWidgets import QApplication, QMainWindow, QWidget, QVBoxLayout, QTextEdit, QPushButton, QComboBox, QLabel, QFileDialog
from PyQt6.QtCore import Qt
from openai import AzureOpenAI
from docx import Document
from PyQt6.QtWidgets import QTabWidget

# COMPUTACIÓN UBICUA Y NUEVOS MODELOS DE APLICACIÓN WEB
# Máster Universitario en Comercio Electrónico | UCM
# Grupo: Wenjun Cai, Jiaxin Qiu, Chenhang He


class TranslatorApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle('Traductor')
        self.setGeometry(100, 100, 400, 600)
        self.target_language = 'English'  # Idioma por defecto al que traducir
        self.initAzureClient()
        self.initUI()
        self.applyStyles()
        
    def applyStyles(self):
        # Global style
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f0f0f0;
            }
            QLabel {
                font-size: 14px;
                color: #333;
            }
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border-radius: 5px;
                padding: 10px;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
            QTextEdit {
                border: 1px solid #ccc;
                border-radius: 5px;
                padding: 5px;
            }
            QComboBox {
                border: 1px solid #ccc;
                border-radius: 5px;
                padding: 5px;
                background-color: white;
            }
            QTabWidget::pane { /* The tab widget frame */
                border-top: 2px solid #ccc;
            }
            
            QTabBar::tab {
                background: #eee;
                border: 1px solid #ccc;
                border-bottom-color: #f0f0f0; /* same as pane color */
                border-top-left-radius: 4px;
                border-top-right-radius: 4px;
                padding: 10px;
            }
            
            QTabBar::tab:selected, QTabBar::tab:hover {
                background: white;
            }
            
        """)
        
    def initUI(self):
        self.tabs = QTabWidget()
        self.tab1 = QWidget()
        self.tab2 = QWidget()

        self.tabs.addTab(self.tab1, "Traducir texto")
        self.tabs.addTab(self.tab2, "Traducir documento")

        self.initTab1()
        self.initTab2()

        self.setCentralWidget(self.tabs)
        
    def initTab1(self):
        layout = QVBoxLayout()

        self.label = QLabel("Traductor de texto basado en la API OpenAI\nVersión demo\n\nCOMPUTACIÓN UBICUA Y NUEVOS MODELOS DE APLICACIÓN WEB\n\nMáster Universitario en Comercio Electrónico | UCM\n\n Grupo:\nWenjun Cai\nJiaxin Qiu\nChenhang He\n")
        layout.addWidget(self.label)
        layout.setAlignment(self.label, Qt.AlignmentFlag.AlignCenter)

        self.inputText = QTextEdit()
        layout.addWidget(self.inputText)

        self.languageComboBox = QComboBox()
        self.languageComboBox.addItem('English', 'en')
        self.languageComboBox.addItem('Spanish', 'es')
        self.languageComboBox.addItem('Chinese', 'zh-Hans')
        self.languageComboBox.currentIndexChanged.connect(self.languageChanged)
        layout.addWidget(self.languageComboBox)

        self.translateButton = QPushButton('Traducir')
        self.translateButton.clicked.connect(self.translateText)
        layout.addWidget(self.translateButton)

        self.outputText = QTextEdit()
        self.outputText.setReadOnly(True)
        layout.addWidget(self.outputText)

        self.tab1.setLayout(layout)

    def initTab2(self):
        layout = QVBoxLayout()
        self.label = QLabel("Traductor de documento basado en la API OpenAI\nVersión demo\n\nCOMPUTACIÓN UBICUA Y NUEVOS MODELOS DE APLICACIÓN WEB\n\nMáster Universitario en Comercio Electrónico | UCM\n\n Grupo:\nWenjun Cai\nJiaxin Qiu\nChenhang He\n")
        layout.addWidget(self.label)
        layout.setAlignment(self.label, Qt.AlignmentFlag.AlignCenter)

        self.languageComboBox = QComboBox()
        self.languageComboBox.addItem('English', 'en')
        self.languageComboBox.addItem('Spanish', 'es')
        self.languageComboBox.addItem('Chinese', 'zh-Hans')
        self.languageComboBox.currentIndexChanged.connect(self.languageChanged)
        layout.addWidget(self.languageComboBox)
        self.docTranslateButton = QPushButton('Seleccionar y Traducir Documento Word')
        self.docTranslateButton.clicked.connect(self.translateWordDocument)
        layout.addWidget(self.docTranslateButton)

        self.docOutputLabel = QLabel("Traducción no iniciada")
        layout.addWidget(self.docOutputLabel)

        self.tab2.setLayout(layout)
    
    # Traducir documento Word    
    def translateWordDocument(self): 
        filePath, _ = QFileDialog.getOpenFileName(self, "Seleccionar documento Word", "", "Word Files (*.docx)")
        if filePath:
            document = Document(filePath)
            full_text = []
            for para in document.paragraphs:
                full_text.append(para.text)
            text_to_translate = "\n".join(full_text)

            # texto del mensaje
            message_text = [
                {"role": "system", "content": "You are a professional, authentic translation engine, only returns translations."},
                {"role": "user", "content": f"Translate the text to {self.target_language.lower()} Language, please do not explain my original text, moreover, when considering the translation results, you should take into account the contextual situation and avoid using words that may cause misunderstandings due to cultural differences. I ask that the result of your translation be accurate, fluent, and reflective of the original meaning. text:{text_to_translate}"}
            ]
            print(message_text)

            try:
                # Enviar solicitud de traducción
                response = self.client.chat.completions.create(
                    model="trans",  # Adjust deployment name
                    messages=message_text,
                    temperature=0.7,
                    max_tokens=800,
                    top_p=0.95,
                    frequency_penalty=0,
                    presence_penalty=0,
                    stop=None
                )
                # Obtener texto traducido
                translated_text = response.choices[0].message.content if response.choices else "Translation failed."

                # Guardar el texto traducido en un nuevo documento de Word
                translated_document = Document()
                translated_document.add_paragraph(translated_text)
                new_file_path = os.path.splitext(filePath)[0] + "_translated.docx"
                translated_document.save(new_file_path)
                self.docOutputLabel.setText(f"Traducción completada: {new_file_path}")
                print("Traducción completada.")
            except Exception as e:
                self.docOutputLabel.setText("Error de traducción, por favor compruebe la llamada a la API.")
                print(f"Traducción fallado {e}")



    def initAzureClient(self):
        self.client = AzureOpenAI(
            azure_endpoint="https://trans-gpt-35.openai.azure.com/",  # Azure OpenAI endpoint
            api_key="",  # Azure OpenAI API key
            api_version="2024-02-15-preview"  # API version
        )
        print("Azure Client Initialized.")

    def languageChanged(self, index):
        self.target_language = self.languageComboBox.itemText(index)
        print(f"Lengua de destino cambiada: {self.target_language}")

    # Traducir Texto
    def translateText(self):
        print("Iniciar la traducción...")
        message_text = [
            {"role": "system", "content": "You are a professional, authentic translation engine, only returns translations."},
            {"role": "user", "content": f"Translate the text to {self.target_language.lower()} Language, please do not explain my original text, moreover, when considering the translation results, you should take into account the contextual situation and avoid using words that may cause misunderstandings due to cultural differences. I ask that the result of your translation be accurate, fluent, and reflective of the original meaning. text:{self.inputText.toPlainText()}"}
        ]
        print(message_text)

        try:
            # Enviar solicitud de traducción
            response = self.client.chat.completions.create(
                model="trans-gpt-35", # Adjust deployment name
                messages=message_text,
                temperature=0.7, #
                max_tokens=800,
                top_p=0.95,
                frequency_penalty=0,
                presence_penalty=0,
                stop=None
            )
            # Obtener texto traducido
            translated_text = response.choices[0].message.content if response.choices else "Translation failed."
            self.outputText.setText(translated_text)
            print("Traducción completada.")
        except Exception as e:
            self.outputText.setText("Error de traducción, por favor compruebe la llamada a la API.")
            print(f"Traducción fallado {e}")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    mainWin = TranslatorApp()
mainWin.show()
sys.exit(app.exec())
